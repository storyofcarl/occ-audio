"""Unit tests for multi-format source loading (.fdx / .docx / .pdf / .txt / .md).
Run: python tests/test_script_source_formats.py"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from occ_audio.script_source import (  # noqa: E402
    Beat, SourceDocument, load_source, normalize_speakers, speaking_characters,
)

_passed = 0
_failed = 0


def check(label: str, condition: bool) -> None:
    global _passed, _failed
    if condition:
        _passed += 1
        print(f"  PASS  {label}")
    else:
        _failed += 1
        print(f"  FAIL  {label}")


_SAMPLE_FDX = """<?xml version="1.0" encoding="UTF-8" standalone="no" ?>
<FinalDraft DocumentType="Script" Template="No" Version="6">
  <Content>
    <Paragraph Type="Scene Heading"><Text>INT. KITCHEN - DAY</Text></Paragraph>
    <Paragraph Type="Action"><Text>Rain hammers the window.</Text></Paragraph>
    <Paragraph Type="Character"><Text>MABEL</Text></Paragraph>
    <Paragraph Type="Parenthetical"><Text>(under her breath)</Text></Paragraph>
    <Paragraph Type="Dialogue"><Text>Where's my damn keys?</Text></Paragraph>
    <Paragraph Type="Character"><Text>Tracee's Daughter (V.O.)</Text></Paragraph>
    <Paragraph Type="Dialogue"><Text>On the counter, Nana.</Text></Paragraph>
    <Paragraph Type="Character"><Text>MABEL</Text></Paragraph>
    <Paragraph Type="Dialogue"><Text>Well ain't that something.</Text></Paragraph>
  </Content>
</FinalDraft>
"""


def test_fdx_parses_structured_types() -> None:
    print("load_source: .fdx uses the real Scene Heading/Character/Dialogue/"
          "Action paragraph types, not heuristics")
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "sample.fdx"
        path.write_text(_SAMPLE_FDX, encoding="utf-8")
        doc = load_source(path)
        check("format is screenplay", doc.format == "screenplay")
        kinds = [b.kind for b in doc.beats]
        check("heading, narration, then dialogue beats in document order",
              kinds == ["heading", "narration", "dialogue", "dialogue", "dialogue"])
        check("Parenthetical paragraph produced NO beat (not spoken verbatim)",
              not any("under her breath" in b.text for b in doc.beats))
        check("Mabel speaks twice", speaking_characters(doc)["Mabel"] == 2)
        check("(V.O.) suffix stripped and casing normalized to one character",
              speaking_characters(doc)["Tracee's Daughter"] == 1)
        check("scene heading text captured",
              doc.beats[0].text == "INT. KITCHEN - DAY")


def test_docx_extracts_paragraphs() -> None:
    print("load_source: .docx paragraphs feed the same prose heuristics as .txt")
    try:
        from docx import Document
    except ImportError:
        print("  SKIP  python-docx not installed")
        return
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "sample.docx"
        document = Document()
        document.add_paragraph("Chapter One")
        document.add_paragraph('Hero said, "Hello there."')
        document.save(str(path))
        doc = load_source(path)
        check("format is prose", doc.format == "prose")
        check("a heading beat was found", any(b.kind == "heading" for b in doc.beats))
        check("dialogue attributed to Hero",
              any(b.kind == "dialogue" and b.speaker == "Hero" for b in doc.beats))


def test_pdf_extracts_text() -> None:
    print("load_source: .pdf text feeds the same prose heuristics as .txt")
    try:
        from pypdf import PdfWriter
    except ImportError:
        print("  SKIP  pypdf not installed")
        return
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "sample.pdf"
        # A minimal blank-page PDF - real text extraction from a generated
        # PDF needs a text layer, which pypdf's writer alone can't add
        # trivially, so this just proves the .pdf branch runs without
        # raising and returns a document with the placeholder empty text
        # collapsed sanely.
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with path.open("wb") as fh:
            writer.write(fh)
        doc = load_source(path)
        check("load_source does not raise on a valid (textless) PDF",
              doc.format in ("prose", "screenplay"))


def test_markdown_heading_detected() -> None:
    print("load_source: a Markdown '#' heading is recognized as a chapter break")
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "sample.md"
        path.write_text(
            "## Scene One\n\nHero said, \"Hi there.\"\n\n"
            "## Scene Two\n\nHero said, \"Bye now.\"\n",
            encoding="utf-8")
        doc = load_source(path)
        headings = [b.text for b in doc.beats if b.kind == "heading"]
        check("both markdown headings recognized (# stripped)",
              headings == ["Scene One", "Scene Two"])


def test_normalize_speakers_merges_spelling_variants() -> None:
    print("normalize_speakers: merges declared name-spelling variants onto "
          "one canonical character, case-insensitive on the alias key")
    doc = SourceDocument(path=Path("x"), format="screenplay", beats=[
        Beat(kind="dialogue", text="a", speaker="Tracee’s Daughter"),  # curly quote
        Beat(kind="dialogue", text="b", speaker="Tracee Daughter"),         # dropped apostrophe
        Beat(kind="dialogue", text="c", speaker="TRACEE'S DAUGHTER"),      # straight quote, caps
        Beat(kind="dialogue", text="d", speaker="Mabel"),                  # untouched
    ])
    aliases = {
        "Tracee’s Daughter": "Tracee's Daughter",
        "Tracee Daughter": "Tracee's Daughter",
        "tracee's daughter": "Tracee's Daughter",   # covers the caps variant (lowercased key)
    }
    normalize_speakers(doc, aliases)
    counts = speaking_characters(doc)
    check("all three variants merged into one canonical speaker",
          counts.get("Tracee's Daughter") == 3)
    check("Mabel is untouched", counts.get("Mabel") == 1)
    check("no leftover variant keys remain", len(counts) == 2)


def test_normalize_speakers_noop_when_no_aliases() -> None:
    print("normalize_speakers: a no-op when aliases is empty")
    doc = SourceDocument(path=Path("x"), format="screenplay", beats=[
        Beat(kind="dialogue", text="a", speaker="Mabel"),
    ])
    result = normalize_speakers(doc, {})
    check("speaker unchanged", result.beats[0].speaker == "Mabel")


if __name__ == "__main__":
    test_fdx_parses_structured_types()
    test_docx_extracts_paragraphs()
    test_pdf_extracts_text()
    test_markdown_heading_detected()
    test_normalize_speakers_merges_spelling_variants()
    test_normalize_speakers_noop_when_no_aliases()
    print(f"\n{_passed} passed, {_failed} failed")
    sys.exit(1 if _failed else 0)

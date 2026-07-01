"""Load a novel / screenplay / script and normalize it into a flat list of
``Beat``s the segmenter can chunk uniformly, regardless of source format.

Supported file types (by extension):

* **.fdx** — Final Draft XML. Parsed directly from its own structured
  ``Scene Heading`` / ``Character`` / ``Dialogue`` / ``Action`` paragraph
  types — no heuristics needed, this is the most reliable input format.
  ``Parenthetical`` paragraphs (acting notes like "(under her breath)") are
  intentionally NOT folded into the dialogue text: doing so would put stage
  direction inside the verbatim quoted string Seed Audio would then read
  aloud, breaking METHODOLOGY.md's verbatim-dialogue rule. They're dropped
  for now.
* **.docx** — Word document. Paragraph text is extracted (via
  ``python-docx``) and then run through the same plain-text heuristics as
  .txt/.md, since .docx carries no reliable screenplay-element typing.
* **.pdf** — text extracted per page (via ``pypdf``), then the same
  plain-text heuristics.
* **.txt / .md / anything else** — read as plain text directly.

Once reduced to plain text, two formats are auto-detected:

* **screenplay** — has ``INT./EXT.`` scene headings. Character cues (a short
  ALL-CAPS line) mark the following block as that character's dialogue;
  everything else is action/narration.
* **prose** — a novel/short story. Chapter/section headings are detected by
  a leading "Chapter"/"Part"/"Scene" line, or a Markdown ``#`` heading.
  Quoted dialogue within a paragraph is split out and attributed to a
  nearby "said Name" / "Name said" pattern when one is found; unattributed
  quotes are flagged (``speaker=None``) rather than guessed — a human/agent
  pass should assign them before casting.

This is a best-effort heuristic pass for non-FDX formats, not NLP-grade
character attribution — per METHODOLOGY.md's "understand the model and the
source material before building or spending," the segmenter's draft output
is meant for human/agent review, not blind trust.
"""
from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path

SCENE_HEADING_RE = re.compile(r"^\s*(INT|EXT|INT\./EXT|I/E)[./\s]", re.IGNORECASE)
CUE_RE = re.compile(r"^[A-Z][A-Z0-9 ._'\-]{1,40}$")
CHAPTER_RE = re.compile(r"^\s*(chapter|part|scene)\s+\S+", re.IGNORECASE)
MD_HEADING_RE = re.compile(r"^#{1,6}\s+(.+)$")
QUOTE_RE = re.compile(r"[\"“]([^\"“”]+)[\"”]")
ATTRIB_RE = re.compile(
    r"([A-Z][a-zA-Z'\-]+)\s+(?:said|asked|replied|shouted|whispered|"
    r"murmured|answered|called|cried|muttered)|"
    r"(?:said|asked|replied|shouted|whispered|murmured|answered|called|"
    r"cried|muttered)\s+([A-Z][a-zA-Z'\-]+)"
)
_PAREN_STRIP_RE = re.compile(r"\s*\([^)]*\)\s*")


@dataclass
class Beat:
    kind: str                      # "heading" | "narration" | "dialogue"
    text: str
    speaker: str | None = None     # set only for kind == "dialogue"


@dataclass
class SourceDocument:
    path: Path
    format: str                    # "screenplay" | "prose"
    beats: list[Beat] = field(default_factory=list)


def _blank_split(text: str) -> list[str]:
    return [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]


def _parse_screenplay(text: str) -> list[Beat]:
    beats: list[Beat] = []
    lines = text.splitlines()
    i = 0
    pending_cue: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal pending_cue, buffer
        block = " ".join(buffer).strip()
        if block:
            if pending_cue:
                beats.append(Beat(kind="dialogue", text=block, speaker=pending_cue))
            else:
                beats.append(Beat(kind="narration", text=block))
        pending_cue = None
        buffer = []

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if SCENE_HEADING_RE.match(stripped):
            flush()
            beats.append(Beat(kind="heading", text=stripped))
            i += 1
            continue
        if CUE_RE.match(stripped) and not SCENE_HEADING_RE.match(stripped):
            flush()
            pending_cue = re.sub(r"\s*\(.*?\)\s*", "", stripped).strip().title()
            i += 1
            continue
        buffer.append(stripped)
        i += 1
    flush()
    return beats


def _attribute(before: str, after: str) -> str | None:
    for chunk in (after[:60], before[-60:]):
        m = ATTRIB_RE.search(chunk)
        if m:
            return (m.group(1) or m.group(2)).strip()
    return None


def _split_paragraph_dialogue(paragraph: str) -> list[Beat]:
    beats: list[Beat] = []
    pos = 0
    for m in QUOTE_RE.finditer(paragraph):
        if m.start() > pos:
            gap = paragraph[pos:m.start()].strip()
            if gap:
                beats.append(Beat(kind="narration", text=gap))
        speaker = _attribute(paragraph[:m.start()], paragraph[m.end():])
        beats.append(Beat(kind="dialogue", text=m.group(1).strip(), speaker=speaker))
        pos = m.end()
    tail = paragraph[pos:].strip()
    if tail:
        beats.append(Beat(kind="narration", text=tail))
    if not beats:
        beats.append(Beat(kind="narration", text=paragraph.strip()))
    return beats


def _parse_prose(text: str) -> list[Beat]:
    beats: list[Beat] = []
    for para in _blank_split(text):
        first_line = para.splitlines()[0].strip()
        md_match = MD_HEADING_RE.match(first_line)
        if md_match:
            beats.append(Beat(kind="heading", text=md_match.group(1).strip()))
            rest = para[len(first_line):].strip()
            if rest:
                beats.extend(_split_paragraph_dialogue(rest))
            continue
        if CHAPTER_RE.match(first_line) and len(first_line) < 60:
            beats.append(Beat(kind="heading", text=first_line))
            rest = para[len(first_line):].strip()
            if rest:
                beats.extend(_split_paragraph_dialogue(rest))
            continue
        beats.extend(_split_paragraph_dialogue(para))
    return beats


# --------------------------------------------------------------------------
# Final Draft (.fdx) — structured parse, no heuristics needed
# --------------------------------------------------------------------------
def _fdx_paragraph_text(para: ET.Element) -> str:
    return "".join(t.text or "" for t in para.findall("Text")).strip()


def _normalize_character_name(raw: str) -> str:
    """"TRACEE'S DAUGHTER" / "Tracee's Daughter" / "TRACEE (V.O.)" all
    normalize to the same key ("Tracee's Daughter") so the same character
    isn't split across differently-cased cues, and parenthetical dialogue
    modifiers (V.O., O.S., CONT'D) don't get treated as a distinct speaker."""
    name = _PAREN_STRIP_RE.sub(" ", raw).strip()
    name = re.sub(r"\s+", " ", name)

    def cap_word(word: str) -> str:
        return word[:1].upper() + word[1:].lower() if word else word

    return " ".join(cap_word(w) for w in name.split(" "))


def _parse_fdx(path: Path) -> list[Beat]:
    root = ET.parse(path).getroot()
    content = root.find("Content")
    beats: list[Beat] = []
    if content is None:
        return beats

    pending_speaker: str | None = None
    for para in content.findall("Paragraph"):
        ptype = para.get("Type", "")
        text = _fdx_paragraph_text(para)
        if not text:
            continue
        if ptype == "Scene Heading":
            beats.append(Beat(kind="heading", text=text))
            pending_speaker = None
        elif ptype == "Character":
            pending_speaker = _normalize_character_name(text)
        elif ptype == "Dialogue":
            beats.append(Beat(kind="dialogue", text=text, speaker=pending_speaker))
        elif ptype == "Action":
            beats.append(Beat(kind="narration", text=text))
            pending_speaker = None
        # "Parenthetical" (acting notes like "(under her breath)") is
        # intentionally skipped — see the module docstring.
    return beats


def _load_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "Reading .docx requires python-docx: pip install python-docx"
        ) from exc
    document = Document(str(path))
    return "\n\n".join(p.text for p in document.paragraphs)


def _load_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "Reading .pdf requires pypdf: pip install pypdf"
        ) from exc
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def detect_format(text: str) -> str:
    return "screenplay" if SCENE_HEADING_RE.search(text) else "prose"


def load_source(path: str | Path) -> SourceDocument:
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix == ".fdx":
        return SourceDocument(path=p, format="screenplay", beats=_parse_fdx(p))
    if suffix == ".docx":
        text = _load_docx_text(p)
    elif suffix == ".pdf":
        text = _load_pdf_text(p)
    else:   # .txt, .md, or anything else — read as plain text
        text = p.read_text(encoding="utf-8")

    fmt = detect_format(text)
    beats = _parse_screenplay(text) if fmt == "screenplay" else _parse_prose(text)
    return SourceDocument(path=p, format=fmt, beats=beats)


def speaking_characters(doc: SourceDocument) -> dict[str, int]:
    """Character name -> number of dialogue beats (a proxy for line count,
    used by casting's ``key`` vs ``all`` threshold and by the segmenter's
    "3 most-speaking characters" tie-break)."""
    counts: dict[str, int] = {}
    for beat in doc.beats:
        if beat.kind == "dialogue" and beat.speaker:
            counts[beat.speaker] = counts.get(beat.speaker, 0) + 1
    return counts
